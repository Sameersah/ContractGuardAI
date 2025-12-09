#!/usr/bin/env python3
"""
Document Generator
Creates properly formatted DOCX and PDF files from text content.
"""

import io
import logging
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX files will be created as plain text.")

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("reportlab not available. PDF files will be created as plain text.")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def create_docx_from_text(content: str, title: Optional[str] = None) -> bytes:
    """
    Create a DOCX file from text content.
    
    Args:
        content: The text content to convert
        title: Optional title for the document
        
    Returns:
        bytes: The DOCX file as bytes
    """
    if not DOCX_AVAILABLE:
        # Fallback: return plain text as bytes
        logger.warning("python-docx not available, returning plain text")
        return content.encode('utf-8')
    
    try:
        # Create a new Document
        doc = Document()
        
        # Add title if provided
        if title:
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            if not para_text.strip():
                continue
            
            # Check if it's a heading (starts with # or is all caps)
            para_text_stripped = para_text.strip()
            if para_text_stripped.startswith('#'):
                # Markdown-style heading
                level = len(para_text_stripped) - len(para_text_stripped.lstrip('#'))
                heading_text = para_text_stripped.lstrip('#').strip()
                doc.add_heading(heading_text, level=min(level, 3))
            elif para_text_stripped.isupper() and len(para_text_stripped) < 100:
                # Likely a heading in all caps
                doc.add_heading(para_text_stripped, level=2)
            else:
                # Regular paragraph
                para = doc.add_paragraph(para_text.strip())
                para_format = para.paragraph_format
                para_format.space_after = Pt(6)
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.read()
        
    except Exception as e:
        logger.error(f"Error creating DOCX: {e}")
        # Fallback to plain text
        return content.encode('utf-8')


def create_pdf_from_text(content: str, title: Optional[str] = None) -> bytes:
    """
    Create a PDF file from text content.
    
    Args:
        content: The text content to convert
        title: Optional title for the document
        
    Returns:
        bytes: The PDF file as bytes
    """
    if not PDF_AVAILABLE:
        # Fallback: return plain text as bytes
        logger.warning("reportlab not available, returning plain text")
        return content.encode('utf-8')
    
    try:
        # Create a BytesIO buffer
        buffer = io.BytesIO()
        
        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=RGBColor(0, 0, 0),
            spaceAfter=12,
            alignment=TA_CENTER
        )
        
        # Add title if provided
        if title:
            elements.append(Paragraph(title, title_style))
            elements.append(Spacer(1, 0.2 * inch))
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            if not para_text.strip():
                elements.append(Spacer(1, 0.1 * inch))
                continue
            
            # Clean up the text
            para_text_clean = para_text.strip().replace('\n', '<br/>')
            
            # Check if it's a heading
            if para_text_clean.startswith('#'):
                level = len(para_text_clean) - len(para_text_clean.lstrip('#'))
                heading_text = para_text_clean.lstrip('#').strip()
                style = styles[f'Heading{min(level, 3)}']
                elements.append(Paragraph(heading_text, style))
                elements.append(Spacer(1, 0.1 * inch))
            elif para_text_clean.isupper() and len(para_text_clean) < 100:
                elements.append(Paragraph(para_text_clean, styles['Heading2']))
                elements.append(Spacer(1, 0.1 * inch))
            else:
                elements.append(Paragraph(para_text_clean, styles['Normal']))
                elements.append(Spacer(1, 0.1 * inch))
        
        # Build PDF
        doc.build(elements)
        
        # Get the PDF bytes
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        return pdf_bytes
        
    except Exception as e:
        logger.error(f"Error creating PDF: {e}")
        # Fallback to plain text
        return content.encode('utf-8')

